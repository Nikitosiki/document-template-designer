from docx import Document
from docx2pdf import convert_to
from tkinter import messagebox

import re
import os
import sys
import time
import pypandoc
import argparse
import pythoncom


# Класс описывающий сущьность строки
# с сохранением статуса использованости.
# Используется в значениях, в словаре
class UseString:
    def __init__(self, value):
        self._value = value
        self._is_use = False

    @property
    def value(self):
        return self._value
    
    @property
    def is_use(self):
        return self._is_use
    
    @is_use.setter
    def is_use(self, is_use):
        self._is_use = is_use


# Замена текста в обьекте параграфа
def replace_text_in_paragraph(element, old_text, new_text):
    text = ""
    for run in element.runs:
        text += run.text
        run.text = ""

    element.runs[0].text = text.replace(old_text, new_text)


# Проходить по внутренней структуре параграфа и подбирает значения со словаря
def find_text_in_paragraph(paragraph, pattern_reg, word_mapping):
    out_matchs = re.finditer(pattern_reg, paragraph.text)  # Пример (.group()): {protocol_number}
    for out_match in out_matchs:
        reg_string = out_match.group()[1:-1]

        # replacement = obj if(isinstance(obj := word_mapping.get(reg_string, ''), str)) else obj.value, obj.is_use = True
        
        if reg_string in word_mapping:
            # Если есть такой ключ в словаре
            replacement = word_mapping[reg_string].value
            word_mapping[reg_string].is_use = True
        else:
            replacement = ''

        replace_text_in_paragraph(paragraph, out_match.group(), replacement)


# Главная ф-ция - заменяет злова в docx
# (а так же проверят на все возможные ошибки при замене)
def modify_docx(document, word_mapping):
    pattern_reg = r"\{[^{}]+\}"

    # Замена текста в абзацах
    for paragraph in document.paragraphs:
        find_text_in_paragraph(paragraph, pattern_reg, word_mapping)

    # Замена текста в таблицах
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    find_text_in_paragraph(paragraph, pattern_reg, word_mapping)

    # Проверяем все ли ключи из словаря были использованы
    unused_keys = [key_obj for key_obj, obj in word_mapping.items() if not obj.is_use]
    if unused_keys:
        str_not_used = ", ".join(unused_keys)
        show_error("These keys were not used: " + str_not_used)


# Удаляет пустые строки в таблицах
def remove_empty_table_rows(doc):
    for table in doc.tables:
        rows_to_delete = []
        for i, row in enumerate(table.rows):
            is_empty = True
            for cell in row.cells:
                if cell.text.strip():
                    is_empty = False
                    break
            if is_empty:
                rows_to_delete.append(row)

        for row in rows_to_delete:
            tbl = row._element.getparent()
            tbl.remove(row._element)


# Парсим строку в dictionary
def string_to_dict(dictionary_str):
    dictionary = {}
    items = dictionary_str.strip('{}').split(', ')
    for item in items:
        key, value = item.split(': ')
        key = key.strip("'")
        value = value.strip("'")
        dictionary[key] = UseString(value)
    return dictionary


# Сохраняет файлы (Конвертирует docx в pdf)
def save_file(doc, output_file, save_docx, save_pdf):
    output_file_name = os.path.splitext(os.path.basename(output_file))[0]   # Извлечение названия файла без расширения
    output_file_directory = os.path.dirname(output_file)   # Извлечение пути к файлу
    output_file_docx = output_file_directory + "\\" + output_file_name + ".docx"
    doc.save(output_file_docx)

    if (save_pdf and not save_docx):
        # Сохраняем в pdf
        # convert(output_file_docx, output_file_pdf)
        convert_to(output_file_docx, output_file_directory)
        if os.path.exists(output_file_docx):
            os.remove(output_file_docx)

    if (not save_pdf and save_docx):
        # Сохраняем в docx
        return

    if (save_pdf and save_docx):
        # Сохраняем в docx и pdf
        # convert(output_file_docx, output_file_pdf)
        convert_to(output_file_docx, output_file_directory)


# Отображение ошибки и завершение работы
def show_error(error_mess):
    messagebox.showerror("Error", error_mess)
    sys.exit()


# Запуск
def main(input_file_docx, output_file, clear_table, save_docx, save_pdf, replacements):
    doc = Document(input_file_docx)

    modify_docx(doc, replacements)
    if (clear_table): remove_empty_table_rows(doc)

    if (not save_pdf and not save_docx):
        # Если не выбран метод сохранения, сохраняем в pdf
        save_pdf = True

    save_file(doc, output_file, save_docx, save_pdf)



if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Replacing keywords in a .docx file, clearing blank lines in tables, and saving to .docx and/or .pdf. \nIf --save_docx or --save_pdf is not selected then --save_pdf is used, you can select both save modes at the same time.")
    parser.add_argument('input_file', type=str, help="Input .docx file with path")
    parser.add_argument('output_file', type=str, help="Path to the output file named")
    parser.add_argument('replacements', type=str, help="Dictionary of replacements as a string")  # Path to the JSON representation of the replacements dictionary
    parser.add_argument('--clear_table', action='store_true', help="Add this flag to clear the table")
    parser.add_argument('--save_docx', action='store_true', help="Add this flag to save to the .docx")
    parser.add_argument('--save_pdf', action='store_true', help="Add this flag to save to the .pdf")

    args = parser.parse_args()
    # print("args.replacements: " + args.replacements)
    try:
        main(args.input_file, args.output_file, args.clear_table, args.save_docx, args.save_pdf, string_to_dict(args.replacements))
    except pythoncom.com_error:
        print(" LibreOffice or Microsoft Office object not found! Requires installation to save a .pdf file.")


# ______________________________________
# ____________Usage example_____________
# ______________________________________

# Console:
# PDFhandler.exe "I:\\NEEEET (XAI)\\projects\\DesignerPDFApp\\test_template-1.docx" "I:\\NEEEET (XAI)\\projects\\DesignerPDFApp\\out_template-1.pdf" --clear_table "{'protocol_number': '234-123-1', 'outside_temp': '23°C', 'calibration_date': '01.07.2023', 'atmospheric_pressure': '78', 'sensor_name': 'датчик (Надлишковий тиск)', 'sensor_model': '2140 мод.ряд САФІР', 'sensor_number': '1', 'sensor_upper_limit': '0 / 250 kPa', 't_pressure_1': 'ывавы', 't_reverse_5': 'fsdfsdfsd fsdf sdf sfsdfdsfsdfd', 'max_inaccuracy': '-125,000 % ', 'performed_position': 'пров.інженер з електроніки', 'performed_name': 'ЖУРАВЕЛЬ Д.Ю.', 'certificate': 'Сертифікат відсутній'}"

# ______________________________________
# ______________________________________
# ______________________________________


# def main_test1():
#     # Засекаем начальное время
#     start_time = time.time()


#     input_file = 'I:\\NEEEET (XAI)\\projects\\DesignerPDFApp\\test_template-1'
#     output_file = 'I:\\NEEEET (XAI)\\projects\\DesignerPDFApp\\out_template-1'
#     replacements = {
#         'protocol_number':      UseString('234-123-1'),
#         'outside_temp':         UseString('23°C'),
#         'calibration_date':     UseString('01.07.2023'),
#         'atmospheric_pressure': UseString('78'),
#         'sensor_name':          UseString('датчик (Надлишковий тиск)'),
#         'sensor_model':         UseString('2140 мод.ряд САФІР'),
#         'sensor_number':        UseString('1'),
#         'sensor_upper_limit':   UseString('0 / 250 kPa'),
#         't_pressure_1':         UseString('ывавы'),
#         't_reverse_5':          UseString('fsdfsdfsd fsdf sdf sfsdfdsfsdfd'),

#         'max_inaccuracy':       UseString('-125,000 % '),
#         'performed_position':   UseString('пров.інженер з електроніки'),
#         'performed_name':       UseString('ЖУРАВЕЛЬ Д.Ю.'),
#         'certificate':          UseString('Сертифікат відсутній')
#         }

#     # main(input_file + ".docx", output_file + ".pdf", True, replacements)

#     print(f"Время заполнения: {time.time() - start_time} секунд")
#     start_time = time.time()


#     doc = Document(input_file + '.docx')
#     modify_docx(doc, replacements)
#     remove_empty_table_rows( doc )
#     doc.save(output_file + '.docx')


#     print(f"Время выполнения функции 1: {time.time() - start_time} секунд")
#     start_time = time.time()

#     convert(output_file + '.docx', output_file + '.pdf')

#     # convert_to_pdf(output_file + '.docx', output_file + '.pdf')


#     print(f"Время выполнения функции 2: {time.time() - start_time} секунд")
#     print(f"Файл {output_file} успешно сохранен.")


# main_test1()

# main(
#     "I:\\NEEEET (XAI)\\projects\\DesignerPDFApp\\test_template-1.docx",
#     "I:\\NEEEET (XAI)\\projects\\DesignerPDFApp\\out_template-1.pdf",
#     True,
#     False,
#     True,
#     string_to_dict("{'protocol_number': '234-123-1', 'outside_temp': '23°C', 'calibration_date': '01.07.2023', 'atmospheric_pressure': '78', 'sensor_name': 'датчик (Надлишковий тиск)', 'sensor_model': '2140 мод.ряд САФІР', 'sensor_number': '1', 'sensor_upper_limit': '0 / 250 kPa', 't_pressure_1': 'ывавы', 't_reverse_5': 'fsdfsdfsd fsdf sdf sfsdfdsfsdfd', 'max_inaccuracy': '-125,000 % ', 'performed_position': 'пров.інженер з електроніки', 'performed_name': 'ЖУРАВЕЛЬ Д.Ю.', 'certificate': 'Сертифікат відсутній'}")
#     )