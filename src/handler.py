from docx import Document
from docx.shared import Inches
from helpers.errormes import*
from helpers.docxtopdf import*

import re
import os


# Замена текста в обьекте параграфа
def replace_text_in_paragraph(element, old_text, new_text):
    text = ""
    for run in element.runs:
        text += run.text
        run.text = ""
    element.runs[0].text = text.replace(old_text, new_text)

# Вставка изображения в обьект параграфа
def replace_image_in_paragraph(element, old_text, new_image_path, new_image_width = 0, new_image_height = 0):
    for run in element.runs:
        run.text = ""

    if new_image_width > 0 and new_image_height > 0:
        element.runs[0].add_picture(new_image_path, width=Inches(new_image_width), height=Inches(new_image_height))
    elif new_image_width > 0 and new_image_height < 0:
        element.runs[0].add_picture(new_image_path, width=Inches(new_image_width))
    elif new_image_width < 0 and new_image_height > 0:
        element.runs[0].add_picture(new_image_path, height=Inches(new_image_height))
    else:
        element.runs[0].add_picture(new_image_path)


# Проходить по внутренней структуре параграфа и подбирает значения со словаря
def find_text_in_paragraph(paragraph, pattern_reg, word_mapping):
    out_matchs = re.finditer(pattern_reg, paragraph.text)  # Пример (.group()): {protocol_number}
    for out_match in out_matchs:
        reg_string = out_match.group()[2:-2]

        # replacement = obj if(isinstance(obj := word_mapping.get(reg_string, ''), str)) else obj.value, obj.is_use = True
        
        if reg_string in word_mapping:
            # Если есть такой ключ в словаре
            replacement = word_mapping[reg_string].value
            is_image = word_mapping[reg_string].is_image
            word_mapping[reg_string].is_use = True
        else:
            replacement = ''
            is_image = False

        if is_image:
            replace_image_in_paragraph(paragraph, out_match.group(), replacement, word_mapping[reg_string].image_width, word_mapping[reg_string].image_height)
        else:
            replace_text_in_paragraph(paragraph, out_match.group(), replacement)

        

# Главная ф-ция - заменяет злова в docx
# (а так же проверят на все возможные ошибки при замене)
def modify_docx(document, word_mapping):
    pattern_reg = r"\{\{[^{}]+\}\}"

    # Замена текста в абзацах
    for paragraph in document.paragraphs:
        find_text_in_paragraph(paragraph, pattern_reg, word_mapping)

    # Замена текста в таблицах
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    find_text_in_paragraph(paragraph, pattern_reg, word_mapping)

    # Проверяем все ли ключи из словаря были использованы
    unused_keys = [key_obj for key_obj, obj in word_mapping.items() if not obj.is_use]
    if unused_keys:
        str_not_used = ", ".join(unused_keys)
        show_error("Error. These keys were not used: " + str_not_used)


# Удаляет пустые строки в таблицах
def remove_empty_table_rows(doc):
    for table in doc.tables:
        rows_to_delete = []
        for i, row in enumerate(table.rows):
            is_empty = True
            for cell in row.cells:
                if cell.text.strip():
                    is_empty = False
                    break
            if is_empty:
                rows_to_delete.append(row)

        for row in rows_to_delete:
            tbl = row._element.getparent()
            tbl.remove(row._element)


# Сохраняет файлы (Конвертирует docx в pdf)
def save_file(doc, output_file, save_docx, save_pdf):
    output_file_name = os.path.splitext(os.path.basename(output_file))[0]   # Извлечение названия файла без расширения
    output_file_directory = os.path.dirname(output_file)   # Извлечение пути к файлу
    output_file_docx = output_file_directory + "\\" + output_file_name + ".docx"
    doc.save(output_file_docx)

    if (save_pdf and not save_docx):
        # Сохраняем в pdf
        convert_to(output_file_docx, output_file_directory)
        if os.path.exists(output_file_docx):
            os.remove(output_file_docx)

    if (not save_pdf and save_docx):
        # Сохраняем в docx
        return

    if (save_pdf and save_docx):
        # Сохраняем в docx и pdf
        convert_to(output_file_docx, output_file_directory)


# Запуск
def main(input_file_docx, output_file, clear_table, save_docx, save_pdf, replacements):
    doc = Document(input_file_docx)

    modify_docx(doc, replacements)
    if (clear_table): remove_empty_table_rows(doc)

    if (not save_pdf and not save_docx):
        # Если не выбран метод сохранения, сохраняем в pdf
        save_pdf = True

    save_file(doc, output_file, save_docx, save_pdf)